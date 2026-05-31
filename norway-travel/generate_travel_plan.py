#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成2026年国庆+中秋 挪威极光旅行计划 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_title(text, size=22, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        run.font.color.rgb = color
    return p


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h


def add_body(text, bold=False, size=11, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(text, bold_prefix="", size=11):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def add_table(headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()
    return table


# ==============================================
# 封面
# ==============================================
for _ in range(3):
    doc.add_paragraph()

add_title("2026年国庆+中秋 挪威极光旅行计划", 26, RGBColor(0x1A, 0x47, 0x8A))

# 装饰线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 30)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
run.font.size = Pt(10)

add_title("成都出发 · 精确执行版", 18, RGBColor(0x33, 0x66, 0x99))
add_title("Tromsø（特罗姆瑟）+ 罗弗敦群岛 · 13日深度游", 14, RGBColor(0x66, 0x66, 0x66))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 30)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("路线：成都 → 奥斯陆 → Tromsø → Svolvær（罗弗敦）→ 回国")
run.font.size = Pt(12)
run.bold = True
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("时间：2026年9月25日（周五）— 10月7日（周三）  |  预算：人均约 ¥18,000")
run.font.size = Pt(11)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
add_body("编制日期：2026年5月  |  本计划基于2025-2026冬季运营数据，出发前请再次确认。", size=10)
add_body("背景：中秋加国庆连在一起（9.25-10.7），只需要请4天年休（9.27-9.30），能凑出完整的13天。", size=10)

doc.add_page_break()

# ==============================================
# 一、核心执行时间线
# ==============================================
add_heading_styled("一、从现在开始的核心执行时间线", level=1)

add_table(
    ["时间节点", "要做的事", "具体操作", "预算/费用"],
    [
        ["2026年5-6月", "检查护照有效期", "确认有效期超过2027年4月15日", "¥120（换发）"],
        ["2026年6月", "买国际机票", "盯Google Flights，设价格提醒", "¥7,000-8,500"],
        ["2026年6月", "下载必备App", "详见App清单", "免费"],
        ["2026年6-7月", "订Tromsø酒店", "Booking / Airbnb 选可免费取消", "¥2,500-3,500/人"],
        ["2026年7月", "订极光团", "Chasing Lights 官网直订", "¥1,500-2,000"],
        ["2026年7-8月", "办申根签证", "提前8周递交签证材料", "签证费+¥1,000"],
        ["2026年8月", "买旅行保险", "覆盖申根区+航班延误", "¥300-500"],
        ["2026年8月", "买电话卡", "淘宝/eSIM 提前买好", "¥50-200"],
        ["2026年8月", "换挪威克朗", "中国银行预约购汇", "NOK 2,000-3,000"],
        ["2026年8-9月", "语言学习", "日常英语+基础挪威语10句", "免费App"],
        ["2026年9月上旬", "买保暖装备", "迪卡侬/优衣库", "¥1,500-2,500"],
        ["2026年9月中旬", "最终确认", "在线值机、打包行李", "-"],
    ]
)

# ==============================================
# 二、假期说明
# ==============================================
add_heading_styled("二、假期结构说明", level=1)

add_table(
    ["日期", "星期", "性质", "安排"],
    [
        ["9月25日（周五）", "", "请假出发", "晚上飞"],
        ["9月26日（周六）﹣9月27日（周日）", "", "中秋假期+周末", "在挪威"],
        ["9月28日（周一）﹣9月30日（周三）", "", "请年休4天", "在挪威"],
        ["10月1日（周四）﹣10月7日（周三）", "", "国庆假期", "在挪威+回程"],
    ]
)

add_body("请假策略：仅需休9月28日-30日（3天年休），加上9月27日如果是工作日再休1天，最多4天，凑出完整13天假期。", bold=True, size=11)

# ==============================================
# 三、买机票
# ==============================================
add_heading_styled("三、什么时候买机票（6月最关键）", level=1)

add_body("什么时候买：2026年6月", bold=True, size=12)
add_body("国庆中秋是出境游最旺的时候，机票一天比一天贵。6月还算便宜，7月开始就涨起来了。", size=10)

add_table(
    ["航线方案", "航司", "参考价（往返）", "推荐度"],
    [
        ["成都→赫尔辛基→奥斯陆→Tromsø ★", "Finnair（芬兰航空）", "¥7,000-8,500", "★★★★★"],
        ["成都→北京/上海→奥斯陆→Tromsø", "国航 / SAS", "¥7,000-8,500", "★★★★"],
        ["成都→伊斯坦布尔→奥斯陆→Tromsø", "土耳其航空", "¥6,500-7,500", "★★★"],
        ["成都→奥斯陆（往返）+ 补飞Tromsø", "Finnair + Norwegian", "¥6,500+1,500", "★★★"],
    ]
)

add_body("★ 首推 Finnair（芬兰航空）", bold=True, size=12)
add_body("成都直飞赫尔辛基大概10个小时，再转一次到奥斯陆再到Tromsø，转机时间最短，整体最顺。")
add_body("★ 购票确认要点", bold=True)
add_bullet("确认包含托运行李（至少23kg）", bold_prefix="✅ ")
add_bullet("转机时间别低于2小时", bold_prefix="✅ ")
add_bullet("最好能免费改签或取消", bold_prefix="✅ ")
add_bullet("国庆旺季票量紧张，看到¥8,000以内的就别等了", bold_prefix="💡 ")

doc.add_page_break()

# ==============================================
# 四、订酒店
# ==============================================
add_heading_styled("四、什么时候订酒店 + 在哪里订（6-7月预订）", level=1)

add_body("订酒店推荐 Booking.com，可以免费取消、到店付款。住小木屋用 Airbnb，罗弗敦那边很多带厨房的。有时候直接去酒店官网订反而不贵，还有会员价或送早餐。")
add_table(
    ["平台", "用途", "优点"],
    [
        ["Booking.com ★首推", "订酒店", "免费取消、到店付款、评分靠谱"],
        ["Airbnb", "订民宿/小木屋", "罗弗敦小木屋首选，有厨房可做饭"],
        ["酒店官网", "直接预订", "有时有会员价或送早餐"],
    ]
)

add_body("▎Tromsø 住宿（住5晚：9月26日-10月1日）", bold=True, size=12)
add_table(
    ["酒店", "类型", "参考价/晚", "推荐理由"],
    [
        ["Scandic Ishavshotel ★", "中高端", "¥1,000-1,300", "港口C位，含早"],
        ["Enter City Hotel", "经济型", "¥600-900", "主街上，含早性价比高"],
        ["Clarion Collection With", "中端", "¥900-1,200", "含晚餐（省钱利器）"],
        ["Airbnb 公寓", "民宿", "¥500-800", "有厨房可做饭"],
    ]
)
add_body("建议住 Tromsø 港口或主街 Storgata 附近。国庆期间房源紧张，6到7月就得订好。", size=10)

add_body("▎罗弗敦住宿（住3晚：10月1日-4日）", bold=True, size=12)
add_table(
    ["住宿", "类型", "参考价/晚", "推荐理由"],
    [
        ["Svolvær 海景小木屋 ★", "Airbnb", "¥700-1,100", "最有北欧感，带厨房"],
        ["Thon Hotel Svolvær", "中端酒店", "¥800-1,000", "含早餐，市中心"],
        ["Anker Brygge", "经典海景房", "¥900-1,200", "港口边渔村风格"],
    ]
)

doc.add_page_break()

# ==============================================
# 五、定极光团
# ==============================================
add_heading_styled("五、什么时候定极光团 + 在哪里定（7月预订）", level=1)

add_body("重要提示：10月初极光季刚开始，团次比11-12月少，提前订！", bold=True, size=11)

add_table(
    ["旅行社", "价格（NOK）", "折合¥", "特点", "预订网址"],
    [
        ["Chasing Lights ★首推", "1,500-2,000", "¥1,000-1,300", "极光摄影团，帮拍照", "chasinglights.no"],
        ["Greenlander", "1,300-1,800", "¥900-1,200", "含保暖服+热饮", "greenlandertravel.no"],
        ["Wandering Owl", "1,400-1,900", "¥950-1,300", "小团+天文讲解", "wanderingowl.no"],
        ["Tromsø Safari", "1,200-1,600", "¥800-1,100", "性价比高", "tromsosafari.no"],
    ]
)

add_body("追光策略：", bold=True, size=12)
add_table(
    ["方案", "报团次数", "成功率", "费用"],
    [
        ["★ 推荐", "2次（Day4 + Day5）", "≈75-80%", "¥2,000-2,600"],
        ["保险方案", "3次（Day4 + Day5 + Day6）", "≈95%", "¥3,000-4,000"],
    ]
)

add_body("10月追光特别提醒：", bold=True)
add_bullet("10月 Tromsø 大概18:30天黑，极光观测窗口在21:00到凌晨1点，夜晚够长")
add_bullet("秋分（9月22日）后极光活跃期就来了，10月初效果不错")
add_bullet("不过10月云多，最好报2到3次团，增加看到的概率")
add_bullet("订之前确认旅行社10月初已经开团了——有些要到11月才营业")

add_body("极光预测App：", bold=True)
add_bullet("My Aurora Forecast — KP指数预测")
add_bullet("yr.no（挪威气象局）— 云量预报最准")
add_bullet("Space Weather Live — 太阳风数据")

# ==============================================
# 六、电话卡
# ==============================================
add_heading_styled("六、办理电话卡（8月买好）", level=1)

add_table(
    ["方案", "类型", "费用", "优点", "缺点"],
    [
        ["★ 淘宝欧洲卡", "实体SIM", "¥50-100", "插卡即用，最稳定", "提前买寄到家"],
        ["eSIM（Airalo）", "虚拟SIM", "¥30-80", "不用换卡", "需手机支持"],
        ["国际漫游", "原号漫游", "¥25-50/天", "保持原号", "贵"],
        ["挪威本地卡", "本地SIM", "NOK 200-300", "信号最好", "到机场找柜台"],
    ]
)

add_body("推荐：淘宝提前买欧洲30天通用卡（20-40GB，约¥60-80），出发前一周下单。", bold=True)

doc.add_page_break()

# ==============================================
# 七、货币兑换
# ==============================================
add_heading_styled("七、货币兑换（8月办理）", level=1)

add_table(
    ["项目", "说明"],
    [
        ["换多少", "建议换 NOK 2,000-3,000（≈¥1,400-2,000）"],
        ["哪里换", "中国银行（提前1-2天手机银行预约）"],
        ["什么时候", "出发前1-2周（2026年8月底）"],
        ["支付方式", "Visa/Mastercard 芯片信用卡为主"],
        ["备用", "带2张不同银行的信用卡"],
        ["不要", "不要在机场换汇（汇率极差）"],
    ]
)
add_body("💡 挪威几乎全面刷卡，现金只是应急。", bold=True)

# ==============================================
# 八、穿着
# ==============================================
add_heading_styled("八、穿着装备（详细清单）", level=1)

add_body("10月初挪威温度在0到10°C之间，雪不大（Tromsø可能根本没雪，山顶才有），但风大，体感挺冷。", bold=True, size=11)
add_body("穿衣诀窍就是叠穿：内层保暖、中间抓绒、外头防风防水。", size=10)

add_body("▎衣物清单", bold=True, size=12)
add_table(
    ["层", "物品", "数量", "说明"],
    [
        ["内层", "HEATTECH 保暖内衣（衣+裤）", "2-3套", "不要穿棉的"],
        ["内层", "美利奴羊毛打底衫", "1件", "追极光时用"],
        ["中层", "抓绒衣 / 羊毛衫", "1-2件", "保暖核心层"],
        ["中层", "薄羽绒", "1件", "灵活穿脱"],
        ["外层", "防风防水冲锋衣", "1件", "比羽绒服更实用（10月雨多）"],
        ["外层", "长款羽绒服", "1件", "晚上的追极光穿"],
        ["下装", "防风防水冲锋裤", "1条", "必须防水防雨"],
        ["下装", "加绒打底裤/保暖裤", "2条", "穿在外裤里面"],
        ["袜子", "美利奴羊毛袜", "3-4双", "中高筒"],
        ["手套", "防水保暖手套", "1副", "触屏功能可选"],
        ["帽子", "毛线帽（盖耳）", "1顶", "防风"],
        ["围巾", "羊毛围巾/Buff", "1条", "保护脸部"],
        ["鞋子", "防水防滑靴 ★", "1双", "Timberland/Sorel/迪卡侬"],
        ["雨具", "轻便雨衣/雨伞", "1件", "10月挪威多雨"],
    ]
)
add_body("跟11月不一样的是，用不着极厚的雪裤和冰爪，防水冲锋衣比长羽绒服更实用——10月下雨比下雪多。", size=10)

add_body("▎其他必备装备", bold=True, size=12)
add_table(
    ["物品", "数量", "原因"],
    [
        ["保温杯（500ml+）", "1个", "追极光时热水救命"],
        ["暖宝宝", "10-20片", "贴手脚、贴手机"],
        ["充电宝（10,000mAh+）", "1个", "寒冷天气掉电极快"],
        ["欧标转换插头（C/F型）", "2个", "挪威两脚圆孔"],
        ["插线板", "1个", "全设备同时充电"],
        ["保湿护肤品+润唇膏", "各1", "北欧干燥"],
        ["防水袋/塑料袋", "若干", "装湿鞋、分装"],
    ]
)

doc.add_page_break()

# ==============================================
# 九、语言
# ==============================================
add_heading_styled("九、出发前学什么语言", level=1)

add_table(
    ["语言", "重要程度", "说明"],
    [
        ["英语 ★★★★★", "必需", "挪威人英语非常好，日常交流足够"],
        ["挪威语 ★★", "非必需但加分", "学10句基础表达，当地人听了会很开心"],
    ]
)

add_body("✦ 建议学的10句挪威语：", bold=True, size=12)
add_table(
    ["中文", "挪威语", "发音参考"],
    [
        ["你好", "Hei", "黑"],
        ["谢谢", "Takk", "塔克"],
        ["不好意思/抱歉", "Unnskyld", "温希尔"],
        ["再见", "Ha det", "哈得"],
        ["是的/好的", "Ja", "呀"],
        ["请", "Vær så snill", "维尔缩斯尼尔"],
        ["多少钱？", "Hvor mye koster det?", "沃尔 迷失 库斯特 得"],
        ["卫生间在哪里？", "Hvor er toalettet?", "沃尔 艾尔 脱瓦列特"],
        ["祝你好运！", "Lykke til!", "吕可 替尔"],
        ["干杯！", "Skål!", "思考尔"],
    ]
)
add_body("工具：Duolingo（多邻国）花20h学基础 + YouTube + Google Translate离线包。")

# ==============================================
# 十、App清单
# ==============================================
add_heading_styled("十、必须下载的App", level=1)

add_table(
    ["分类", "App", "用途"],
    [
        ["✈️ 航班", "Google Flights", "查航班、价格提醒"],
        ["✈️ 航班", "航旅纵横 / FlightRadar24", "值机、实时动态"],
        ["🏨 住宿", "Booking.com", "订酒店、管理订单"],
        ["🏨 住宿", "Airbnb", "订小木屋"],
        ["🌌 极光", "My Aurora Forecast", "KP指数预测"],
        ["🌌 极光", "Space Weather Live", "太阳风数据"],
        ["🌤 天气", "yr.no", "挪威气象局云图"],
        ["🌤 天气", "Windy", "云层覆盖预报"],
        ["🗺 导航", "Google Maps", "离线下载地图"],
        ["🗺 导航", "Maps.me", "离线导航省流量"],
        ["🚌 交通", "Tromsø Billett", "Tromsø公交票"],
        ["💬 翻译", "Google Translate", "挪威语离线包"],
        ["💰 汇率", "极简汇率", "实时换算"],
        ["📸 拍照", "Lightroom / Snapseed", "极光后期调色"],
        ["📞 通讯", "Airalo（eSIM）", "流量套餐"],
        ["📞 通讯", "WhatsApp", "联系旅行社"],
        ["🍽 美食", "TripAdvisor", "查餐厅评价"],
    ]
)

doc.add_page_break()

# ==============================================
# 十一、预算
# ==============================================
add_heading_styled("十一、预算总表（人均）", level=1)

add_table(
    ["项目", "预算（¥）", "说明"],
    [
        ["国际机票（成都↔Tromsø）", "7,500-8,500", "国庆旺季价格偏高"],
        ["挪威境内交通", "2,000-2,500", "飞机+渡轮/巴士"],
        ["住宿（8晚分摊后）", "3,000-4,000", "两人分摊，含早为主"],
        ["极光团（2次）", "1,500-2,000", "Chasing Lights"],
        ["餐饮（13天）", "2,000-2,500", "含早省一餐，日均¥180"],
        ["市内交通+景点", "500-1,000", "公交、缆车、博物馆"],
        ["签证+保险", "1,000-1,500", "申根签证+旅行保险"],
        ["装备购买", "1,500-2,500", "从零买（可重复用）"],
        ["", "", ""],
        ["★ 总计（不含装备）", "17,500-22,000", "建议心理准备 ¥20,000"],
        ["★ 总计（含装备）", "19,000-24,500", "如果从零买齐"],
    ]
)
add_body("国庆算旺季，机票比11月贵个一千左右，但白天长不少（7:30天亮，18:00才黑），能玩的时间更多。", size=10)

# ==============================================
# 十二、逐日行程
# ==============================================
add_heading_styled("十二、逐日行程表（13天）", level=1)

add_table(
    ["天", "日期", "星期", "活动概要", "住宿"],
    [
        ["Day1", "9/25", "五", "成都出发，晚上飞", "飞机上"],
        ["Day2", "9/26", "六", "抵达Tromsø，超市采购，港口散步，倒时差", "Tromsø"],
        ["Day3", "9/27", "日", "中秋🎑 北极教堂→缆车→极地博物馆", "Tromsø"],
        ["Day4", "9/28", "一", "市区自由探索→20:00 第一次追极光✨", "Tromsø"],
        ["Day5", "9/29", "二", "狗拉雪橇/峡湾游→第二次追极光✨", "Tromsø"],
        ["Day6", "9/30", "三", "自由日/补追极光✨", "Tromsø"],
        ["Day7", "10/1", "四", "国庆🎉 飞Svolvær（罗弗敦）", "罗弗敦"],
        ["Day8", "10/2", "五", "罗弗敦：渔村漫游、拍照", "罗弗敦"],
        ["Day9", "10/3", "六", "罗弗敦：登山徒步/出海", "罗弗敦"],
        ["Day10", "10/4", "日", "罗弗敦→奥斯陆", "奥斯陆/飞机"],
        ["Day11", "10/5", "一", "奥斯陆市区游（可选）→飞回程", "飞机上"],
        ["Day12", "10/6", "二", "转机中", "飞机上"],
        ["Day13", "10/7", "三", "抵达成都 🏠", "到家"],
    ]
)

doc.add_page_break()

# ==============================================
# 十三、关键提醒
# ==============================================
add_heading_styled("十三、关键提醒（出行前必读）", level=1)

add_body("✦ 极光", bold=True, size=13)
add_bullet("10月初极光季刚开始，团次可能较少，7月就预订")
add_bullet("极光活跃期21:00-01:00，要熬夜跟团")
add_bullet("2-3次团+好心态=最佳策略")

add_body("✦ 天气", bold=True, size=13)
add_bullet("10月挪威温度0-10°C，比11月暖和但雨水多")
add_bullet("防水冲锋衣比羽绒服更重要")
add_bullet("可能会有雨夹雪，路面湿滑注意安全")

add_body("挪威治安不错，Tromsø 在极光城市里也算安全的。", bold=True, size=13)
add_bullet("紧急电话：112（报警）113（急救）")
add_bullet("中国驻挪威大使馆：+47 22 59 48 00")

add_body("✦ 省钱技巧", bold=True, size=13)
add_bullet("住酒店选带早餐的，省一顿饭钱")
add_bullet("水和零食去超市买，便利店贵一半以上")
add_bullet("随身带个保温杯装热水")
add_bullet("购物满 NOK 315 可以退税（12%-19%），记得办")

doc.add_page_break()

# ==============================================
# 十四、罗弗敦深度游玩指南
# ==============================================
add_heading_styled("十四、罗弗敦群岛深度游玩指南", level=1)

add_body("罗弗敦才是这趟旅程最出彩的地方。从 Tromsø 飞到 Svolvær 只要1小时，但风景完全不一样——峡湾、渔村、雪山交替出现，怎么拍都好看。", size=11)

add_body("▎罗弗敦必做清单", bold=True, size=12)

add_table(
    ["活动", "地点", "时间", "推荐理由"],
    [
        ["E10 公路自驾", "全岛", "半天", "号称\"最美公路\"，一边靠山一边临海"],
        ["Reinebringen 徒步", "Reine", "2-3小时", "山顶俯瞰罗弗敦经典画面，有点陡但值得"],
        ["Å 渔村探访", "Å i Lofoten", "半日", "世界尽头的小渔村，鳕鱼干架子是标志"],
        ["Henningsvær 足球场", "Henningsvær", "1小时", "建在礁石上的足球场，很上镜"],
        ["出海钓鱼", "Svolvær 出发", "3-4小时", "挪威海钓，钓上来直接船上煮"],
        ["极光拍摄", "天气好随处可拍", "夜间", "罗弗敦的极光配峡湾，比Tromsø更出片"],
        ["鳕鱼干博物馆", "Å", "1小时", "了解罗弗敦百年渔业历史"],
    ]
)

add_body("▎罗弗敦怎么玩", bold=True, size=12)
add_body("主路 E10 公路把各个渔村串在一起。租车最自由，从 Svolvær 开到 Å 单程约2小时，边走边停。也可以坐巴士，但班次不多。最省心的是到了当地报一日游团。")
add_body("10月初是淡旺季交替期，游客少、住宿不贵，但部分餐厅可能关门，出发前查一下营业时间。", size=10)

doc.add_page_break()

# ==============================================
# 十五、挪威美食推荐
# ==============================================
add_heading_styled("十五、挪威吃什么（餐厅推荐）", level=1)

add_body("挪威不算美食大国，但在 Tromsø 和罗弗敦还是能找到不错的馆子。", size=11)

add_body("Tromsø", bold=True, size=12)
add_table(
    ["餐厅", "类型", "人均（NOK）", "推荐菜"],
    [
        ["Fiskekompaniet ★", "海鲜", "300-500", "鳕鱼舌、三文鱼，当地口碑靠前"],
        ["Raketten", "街头快餐", "50-150", "热狗加鱼子酱——Tromsø特色小吃"],
        ["Hildr Gastrobar", "创意料理", "400-600", "驯鹿肉塔塔、北极甜品"],
        ["Supremeon", "越南粉", "150-250", "冷天来一碗很暖"],
        ["Burgr", "汉堡", "200-300", "本地牛肉汉堡连锁"],
        ["Yonas Pizzeria", "披萨", "150-250", "Tromsø老牌披萨店，深夜还营业"],
    ]
)

add_body("罗弗敦", bold=True, size=12)
add_table(
    ["餐厅", "位置", "人均（NOK）", "推荐菜"],
    [
        ["Måltid ★", "Svolvær", "300-500", "本地食材创新菜，鳕鱼做得特别好"],
        ["Bacalao", "Svolvær", "200-350", "地中海风的挪威海鲜"],
        ["Anitas Seafood", "Sakrisøy", "100-200", "鱼饼汉堡必吃，路过别错过"],
        ["Krambua", "Reine", "150-250", "渔村小馆，鱼汤很鲜"],
        ["Lofoten Fiskerestaurant", "Svolvær", "400-600", "港口边，吃风景也吃味道"],
    ]
)

add_body("省钱吃喝建议：", bold=True, size=11)
add_bullet("早餐靠酒店，住含早的能省不少")
add_bullet("午餐在超市买材料自己做（小木屋都有厨房）")
add_bullet("晚饭出去吃，不用顿顿大餐，搭配着来")
add_bullet("超市推荐 Rema 1000、Kiwi、Coop，比便利店便宜很多")
add_bullet("挪威自来水可以直接喝，不用买瓶装水")

doc.add_page_break()

# ==============================================
# 十六、极光拍照指南
# ==============================================
add_heading_styled("十六、极光拍照指南（手机+相机）", level=1)

add_body("花了几千块追极光，拍不出来就太可惜了。这部分不管用相机还是手机都能跟上。", size=11)

add_body("相机设置", bold=True, size=12)
add_table(
    ["参数", "推荐值", "说明"],
    [
        ["光圈", "f/1.8 - f/2.8", "越大越好，进光量是核心"],
        ["快门速度", "3-15秒", "极光弱用15秒，爆发时3-5秒就行"],
        ["ISO", "800-3200", "先试800，不够再往上加"],
        ["对焦", "手动无限远", "拧到∞标志再回调一点点"],
        ["白平衡", "3500-4000K", "偏冷色调极光更绿"],
        ["RAW格式", "开启", "后期调色空间大很多"],
    ]
)
add_body("必备工具：三脚架（必须）、快门线或定时拍摄、备用电池（低温掉电极快）。", size=10)
add_body("新手最容易翻车的是对焦。自动对焦在夜里基本废的，一定要切手动。出发前在家练几次。", size=10)

add_body("手机也能拍", bold=True, size=12)
add_table(
    ["手机", "方法"],
    [
        ["iPhone 12及以上", "夜间模式 → 曝光拉到10秒 → 放稳（最好用三脚架）"],
        ["华为/Pixel", "夜景模式效果不错，同样要稳"],
        ["安卓通用", "Pro/专业模式 → ISO 1600 → 快门10-15秒 → 手动对焦无限远"],
        ["推荐App", "NightCap（iPhone）/ Camera FV-5（安卓），能手动控制全部参数"],
    ]
)
add_body("核⼼就三条：稳（三脚架）、长曝光（几秒到十几秒）、手动对焦。出发前摸黑熟悉一下操作，比现场手忙脚乱强。", size=10)

doc.add_page_break()

# ==============================================
# 十七、挪威境内交通衔接
# ==============================================
add_heading_styled("十七、挪威境内交通怎么串", level=1)

add_body("这趟行程涉及三段交通：成都→Tromsø、Tromsø→罗弗敦、罗弗敦→奥斯陆→回国。", size=11)

add_body("▎Tromsø → 罗弗敦（Svolvær）", bold=True, size=12)
add_table(
    ["方式", "时长", "价格", "备注"],
    [
        ["★ 飞机（最推荐）", "1小时", "¥500-800", "Norwegian/Widerøe 执飞，提前买票"],
        ["渡轮 Hurtigruten", "12小时", "¥600-1,000", "沿海航线，可省一晚住宿"],
        ["巴士（需转车）", "8-10小时", "¥300-500", "最便宜但太折腾，不推荐"],
    ]
)
add_body("飞机虽然贵但省时间。渡轮如果时间对得上值得一试——沿着挪威海岸线走，本身就是风景。", size=10)

add_body("▎罗弗敦当地交通", bold=True, size=12)
add_bullet("租车：最自由，Svolvær 机场取还，约¥400-600/天，10月初路况好，不需要雪胎")
add_bullet("巴士：班次少，周末部分线路停运，提前查时刻表")
add_bullet("出租车：贵，只适合短途应急")
add_body("如果会开车，罗弗敦值得租车。E10 公路风景太好，随走随停，巴士做不到。", size=10)

add_body("▎罗弗敦 → 奥斯陆 → 成都", bold=True, size=12)
add_bullet("Svolvær/Bodø → Oslo：Norwegian/SAS 约2小时，¥500-1,000")
add_bullet("Oslo → 成都：回程看买的票，通常也要转一次机")
add_bullet("建议留够中转时间（国际段≥3h，境内段≥1.5h）")

doc.add_page_break()

# ==============================================
# 十八、应急处理方案
# ==============================================
add_heading_styled("十八、突发情况怎么办", level=1)

add_body("提前想想应对方案，真遇到了不至于慌。", size=11)

add_table(
    ["情况", "应对方法"],
    [
        ["航班取消或延误", "马上联系航司改签，用 FlightRadar24 提前关注前序航班状态"],
        ["错过中转", "找航司柜台安排下一班，芬兰航空或国航一般免费改签"],
        ["极光团取消", "旅行社一般提前几小时通知，约好后备日期"],
        ["天气恶劣", "关注 yr.no 和 Windy 云图，10月挪威暴风雨不算多但遇上别出门"],
        ["生病或受伤", "带常用药（感冒、肠胃、止痛），严重打113叫救护车"],
        ["被盗或丢失物品", "挪威治安不错但别大意，随身包拉好，丢东西找警察（112）"],
        ["护照丢失", "联系中国驻挪威大使馆（+47 22 59 48 00）办旅行证"],
        ["信用卡或现金丢失", "带2张不同银行的卡分开放，现金少量应急就行"],
    ]
)

add_body("随身必带的应急物品：", bold=True, size=11)
add_bullet("护照复印件（和原件分开放）")
add_bullet("电子版：护照、签证、机票、酒店订单存手机和云端")
add_bullet("常用药：感冒药、退烧药、肠胃药、止痛药、创可贴")
add_bullet("充电宝（10,000mAh以上，低温掉电快）")
add_bullet("紧急联系人信息写纸上放钱包")

add_body("保险提醒：", bold=True, size=11)
add_body("买旅行保险确认覆盖医疗运送（50万元以上）、航班延误（每4小时赔500-1,000元）、行李丢失或延误。推荐安联或美亚的申根险，200-400元搞定。", size=10)

doc.add_page_break()

# ==============================================
# 十九、购物与退税指南
# ==============================================
add_heading_styled("十九、买什么 + 怎么退税", level=1)

add_body("挪威物价不便宜，但有些东西值得带回来。", size=11)

add_body("▎买什么", bold=True, size=12)
add_table(
    ["商品", "推荐理由", "参考价（NOK）"],
    [
        ["挪威羊毛衫 Lusekofte", "经典民族图案，保暖又耐穿", "800-2,000"],
        ["极光周边小物", "Tromsø纪念品店有很多极光主题的", "100-500"],
        ["三文鱼或鱼子酱", "真空包装可以带上飞机", "200-500"],
        ["驯鹿皮", "当地人家几乎都有，手感好", "1,000-2,500"],
        ["挪威巧克力", "Freia 和 Nidar 国民品牌，送人合适", "30-80/板"],
        ["户外装备", "挪威本土品牌 Norrøna、Bergans 比国内便宜", "看具体"],
        ["维京周边", "船模型、符文饰品，有意思的小纪念品", "50-300"],
    ]
)

add_body("▎退税流程", bold=True, size=12)
add_body("挪威不是欧盟，但有自己的退税政策：")
add_bullet("最低消费：单店满 NOK 315（约¥200）")
add_bullet("退税率：12%-19%，金额越高退越多")
add_bullet("购物时出示护照，店员开退税单 → 离境时在机场 Tax Free 柜台办理 → 退到信用卡或拿现金")
add_bullet("Tromsø 机场和奥斯陆机场（Gardermoen）都有退税点")
add_body("注意：退税在离开挪威的最后一站统一办。如果路线是 Tromsø → Oslo → 回国，就在 Oslo 机场办。", bold=True, size=10)

doc.add_page_break()

# ==============================================
# 附录：快速参照表
# ==============================================
add_heading_styled("附录：关键信息速查", level=1)

add_table(
    ["项目", "内容"],
    [
        ["出发日期", "9月25日（周五）出发 → 10月7日（周三）到成都"],
        ["请假策略", "只需请9月28-30日（3-4天年休）"],
        ["国际航班", "推荐 Finnair 成都→赫尔辛基→奥斯陆→Tromsø"],
        ["签证", "申根签证，提前8周递交（7-8月办）"],
        ["住宿", "Tromsø 5晚 + 罗弗敦 3晚（6-7月预订）"],
        ["极光团", "报2-3次团，7月预订 Chasing Lights"],
        ["总预算", "人均约 ¥18,000-22,000（不含装备）"],
        ["核心提醒", "防水冲锋衣 > 羽绒服；超市比便利店便宜一半"],
    ]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 30)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("★ 一句话总结 ★")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

add_body("6月买机票 → 6到7月订酒店和极光团 → 7到8月办签证、买电话卡、换钱 → 9月准备装备 → 9月25日出发！", bold=True, size=12)

add_body("中秋在北极看极光，国庆在罗弗敦看海，挺会安排的。", size=11)

# ==============================================
# 保存
# ==============================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "2026年国庆+中秋_挪威极光旅行计划_详细版.docx")
# 如果文件被占用，先删旧的再写
if os.path.exists(output_path):
    try:
        os.remove(output_path)
    except PermissionError:
        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                   "2026年国庆+中秋_挪威极光旅行计划_优化版.docx")
doc.save(output_path)
print(f"已生成: {output_path}")
print(f"文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
